"""
DOC文档导出模块

将生成的书籍内容导出为Word文档，包含目录、样式设置等功能。
"""

import os
import re
from typing import Optional, List, Tuple
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

from .outline_generator import BookOutline
from .config import get_config


class DocExporter:
    """DOC文档导出器
    
    将书籍内容导出为格式化的Word文档，支持：
    - 自动生成目录
    - 自定义样式设置
    - 章节标题层级
    - 页眉页脚
    
    Attributes:
        config: 配置对象
        document: Word文档对象
    
    Example:
        >>> exporter = DocExporter()
        >>> exporter.create_document(outline, content_dict)
        >>> exporter.save("output.docx")
    """
    
    def __init__(self) -> None:
        """初始化DOC导出器"""
        self.config = get_config()
        self.document: Optional[Document] = None
        self.doc_config = self.config.get_document_config()
    
    def create_document(
        self,
        outline: BookOutline,
        content_dict: dict
    ) -> Document:
        """创建Word文档
        
        Args:
            outline: 书籍大纲
            content_dict: 内容字典，键为章节号（0表示自序），值为内容
            
        Returns:
            创建的Word文档对象
        """
        self.document = Document()
        
        # 设置文档默认字体
        self._setup_document_styles()
        
        # 添加封面
        self._add_cover_page(outline)
        
        # 添加目录
        self._add_table_of_contents(outline, content_dict)
        
        # 添加自序
        if 0 in content_dict:
            self._add_preface(content_dict[0])
        
        # 添加各章节
        for chapter in outline.chapters:
            if chapter.chapter_number in content_dict:
                self._add_chapter(chapter, content_dict[chapter.chapter_number])
        
        return self.document
    
    def _setup_document_styles(self) -> None:
        """设置文档样式"""
        if not self.document:
            return
        
        # 获取样式配置
        body_font = self.doc_config.get('body_font', '宋体')
        body_size = self.doc_config.get('body_size', 12)
        title_font = self.doc_config.get('title_font', '黑体')
        line_spacing = self.doc_config.get('line_spacing', 1.5)
        
        # 设置正文样式
        style = self.document.styles['Normal']
        style.font.name = body_font
        style.font.size = Pt(body_size)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), body_font)
        style.paragraph_format.line_spacing = line_spacing
        style.paragraph_format.space_after = Pt(6)
        
        # 设置标题1样式（章标题）
        heading1 = self.document.styles['Heading 1']
        heading1.font.name = title_font
        heading1.font.size = Pt(18)
        heading1.font.bold = True
        heading1._element.rPr.rFonts.set(qn('w:eastAsia'), title_font)
        heading1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        heading1.paragraph_format.space_before = Pt(24)
        heading1.paragraph_format.space_after = Pt(18)
        
        # 设置标题2样式（节标题）
        heading2 = self.document.styles['Heading 2']
        heading2.font.name = title_font
        heading2.font.size = Pt(14)
        heading2.font.bold = True
        heading2._element.rPr.rFonts.set(qn('w:eastAsia'), title_font)
        heading2.paragraph_format.space_before = Pt(18)
        heading2.paragraph_format.space_after = Pt(12)
        
        # 设置标题3样式（小节标题）
        heading3 = self.document.styles['Heading 3']
        heading3.font.name = title_font
        heading3.font.size = Pt(12)
        heading3.font.bold = True
        heading3._element.rPr.rFonts.set(qn('w:eastAsia'), title_font)
        heading3.paragraph_format.space_before = Pt(12)
        heading3.paragraph_format.space_after = Pt(6)
    
    def _add_cover_page(self, outline: BookOutline) -> None:
        """添加封面页
        
        Args:
            outline: 书籍大纲
        """
        if not self.document:
            return
        
        title_font = self.doc_config.get('title_font', '黑体')
        
        # 添加空行
        for _ in range(6):
            self.document.add_paragraph()
        
        # 主标题
        title_para = self.document.add_paragraph()
        title_run = title_para.add_run(outline.title)
        title_run.font.name = title_font
        title_run.font.size = Pt(26)
        title_run.font.bold = True
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), title_font)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 副标题
        if outline.subtitle:
            self.document.add_paragraph()
            subtitle_para = self.document.add_paragraph()
            subtitle_run = subtitle_para.add_run(outline.subtitle)
            subtitle_run.font.name = title_font
            subtitle_run.font.size = Pt(16)
            subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), title_font)
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 分页
        self.document.add_page_break()
    
    def _add_table_of_contents(
        self,
        outline: BookOutline,
        content_dict: dict
    ) -> None:
        """添加目录页
        
        Args:
            outline: 书籍大纲
            content_dict: 内容字典
        """
        if not self.document:
            return
        
        title_font = self.doc_config.get('title_font', '黑体')
        
        # 目录标题
        toc_title = self.document.add_paragraph()
        toc_run = toc_title.add_run("目  录")
        toc_run.font.name = title_font
        toc_run.font.size = Pt(18)
        toc_run.font.bold = True
        toc_run._element.rPr.rFonts.set(qn('w:eastAsia'), title_font)
        toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        toc_title.paragraph_format.space_after = Pt(24)
        
        # 自序
        if 0 in content_dict:
            self._add_toc_entry("自序", 0)
        
        # 章节条目
        for chapter in outline.chapters:
            if chapter.chapter_number in content_dict:
                entry_text = f"第{chapter.chapter_number}章  {chapter.title}"
                self._add_toc_entry(entry_text, 0)
        
        # 分页
        self.document.add_page_break()
    
    def _add_toc_entry(self, text: str, level: int) -> None:
        """添加目录条目
        
        Args:
            text: 条目文本
            level: 层级（0=一级，1=二级）
        """
        if not self.document:
            return
        
        body_font = self.doc_config.get('body_font', '宋体')
        
        para = self.document.add_paragraph()
        
        # 添加缩进
        if level > 0:
            para.paragraph_format.left_indent = Inches(0.5 * level)
        
        run = para.add_run(text)
        run.font.name = body_font
        run.font.size = Pt(12)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), body_font)
    
    def _add_preface(self, content: str) -> None:
        """添加自序
        
        Args:
            content: 自序内容
        """
        if not self.document:
            return
        
        title_font = self.doc_config.get('title_font', '黑体')
        body_font = self.doc_config.get('body_font', '宋体')
        
        # 标题
        title_para = self.document.add_paragraph()
        title_run = title_para.add_run("自序")
        title_run.font.name = title_font
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), title_font)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_after = Pt(18)
        
        # 内容
        self._add_formatted_content(content)
        
        # 分页
        self.document.add_page_break()
    
    def _add_chapter(self, chapter, content: str) -> None:
        """添加章节内容
        
        Args:
            chapter: 章节大纲对象
            content: 章节内容
        """
        if not self.document:
            return
        
        title_font = self.doc_config.get('title_font', '黑体')
        
        # 章节标题
        title_para = self.document.add_paragraph()
        title_run = title_para.add_run(f"第{chapter.chapter_number}章  {chapter.title}")
        title_run.font.name = title_font
        title_run.font.size = Pt(18)
        title_run.font.bold = True
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), title_font)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_after = Pt(18)
        
        # 内容
        self._add_formatted_content(content)
        
        # 分页（除最后一章外）
        self.document.add_page_break()
    
    def _add_formatted_content(self, content: str) -> None:
        """添加格式化内容
        
        解析内容中的标题和段落，应用相应样式。
        
        Args:
            content: 文本内容
        """
        if not self.document:
            return
        
        body_font = self.doc_config.get('body_font', '宋体')
        title_font = self.doc_config.get('title_font', '黑体')
        
        lines = content.split('\n')
        
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            # 检测是否是标题行
            if self._is_section_title(line):
                # 作为小节标题
                para = self.document.add_paragraph()
                run = para.add_run(line)
                run.font.name = title_font
                run.font.size = Pt(14)
                run.font.bold = True
                run._element.rPr.rFonts.set(qn('w:eastAsia'), title_font)
                para.paragraph_format.space_before = Pt(12)
                para.paragraph_format.space_after = Pt(6)
            elif self._is_subsection_title(line):
                # 作为小标题
                para = self.document.add_paragraph()
                run = para.add_run(line)
                run.font.name = title_font
                run.font.size = Pt(12)
                run.font.bold = True
                run._element.rPr.rFonts.set(qn('w:eastAsia'), title_font)
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(3)
            else:
                # 普通段落
                para = self.document.add_paragraph()
                run = para.add_run(line)
                run.font.name = body_font
                run.font.size = Pt(12)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), body_font)
                para.paragraph_format.first_line_indent = Inches(0.5)
                para.paragraph_format.line_spacing = 1.5
    
    def _is_section_title(self, line: str) -> bool:
        """检查是否是小节标题
        
        Args:
            line: 文本行
            
        Returns:
            是否是小节标题
        """
        # 常见的节标题模式
        patterns = [
            r'^第[一二三四五六七八九十\d]+节',
            r'^\d+\.\d+',
            r'^[（(]\d+[)）]',
        ]
        
        for pattern in patterns:
            if re.match(pattern, line):
                return True
        
        # 短句且没有标点符号，可能是标题
        if len(line) < 20 and not re.search(r'[。，；：]', line):
            return True
        
        return False
    
    def _is_subsection_title(self, line: str) -> bool:
        """检查是否是小标题
        
        Args:
            line: 文本行
            
        Returns:
            是否是小标题
        """
        patterns = [
            r'^\d+\.\d+\.\d+',
            r'^[（(][一二三四五六七八九十][)）]',
        ]
        
        for pattern in patterns:
            if re.match(pattern, line):
                return True
        
        return False
    
    def save(self, filepath: str) -> str:
        """保存文档
        
        Args:
            filepath: 保存路径
            
        Returns:
            保存的文件路径
        """
        if not self.document:
            raise ValueError("文档尚未创建")
        
        # 确保目录存在
        os.makedirs(os.path.dirname(filepath) if os.path.dirname(filepath) else '.', exist_ok=True)
        
        self.document.save(filepath)
        return filepath


def export_to_docx(
    outline: BookOutline,
    content_dict: dict,
    output_path: str
) -> str:
    """便捷函数：导出为DOCX文件
    
    Args:
        outline: 书籍大纲
        content_dict: 内容字典
        output_path: 输出路径
        
    Returns:
        输出文件路径
    """
    exporter = DocExporter()
    exporter.create_document(outline, content_dict)
    return exporter.save(output_path)
